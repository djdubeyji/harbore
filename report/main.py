"""
Harbore Report Engine
Generates Word (.docx) and PDF security assessment reports from scan findings.
"""
import io
import os
import logging
from datetime import datetime
from typing import Optional
from fastapi import FastAPI, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.platypus import PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

load_dotenv()
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
log = logging.getLogger(__name__)

# ─── Data models ─────────────────────────────────────────────────────────────

class FindingModel(BaseModel):
    title: str
    severity: str
    module: str
    cvss_score: Optional[float] = None
    endpoint: Optional[str] = None
    method: Optional[str] = None
    description: str
    owasp_ref: Optional[str] = None
    pci_requirement: Optional[str] = None
    cwe_id: Optional[str] = None
    ai_summary: Optional[str] = None
    ai_remediation: Optional[str] = None
    request: Optional[str] = None
    response: Optional[str] = None

class FailureModel(BaseModel):
    target: str
    module: str
    attempts: int
    final_error: str

class ReportRequest(BaseModel):
    scan_id: str
    scan_name: str
    target_count: int
    modules_run: list[str]
    duration_mins: Optional[float] = None
    stats: dict[str, int] = {}
    findings: list[FindingModel] = []
    failures: list[FailureModel] = []
    executive_summary: Optional[str] = None
    critical_section: Optional[str] = None
    remediation_priorities: Optional[str] = None
    pci_narrative: Optional[str] = None
    format: str = "docx"  # "docx" | "pdf"

# ─── Severity config ──────────────────────────────────────────────────────────

SEV_ORDER = ["critical", "high", "medium", "low", "info"]
SEV_COLORS_RGB = {
    "critical": RGBColor(0xEF, 0x44, 0x44),
    "high":     RGBColor(0xF9, 0x73, 0x16),
    "medium":   RGBColor(0xEA, 0xB3, 0x08),
    "low":      RGBColor(0x3B, 0x82, 0xF6),
    "info":     RGBColor(0x6B, 0x72, 0x80),
}
SEV_COLORS_PDF = {
    "critical": colors.HexColor("#EF4444"),
    "high":     colors.HexColor("#F97316"),
    "medium":   colors.HexColor("#EAB308"),
    "low":      colors.HexColor("#3B82F6"),
    "info":     colors.HexColor("#6B7280"),
}

# ─── FastAPI app ──────────────────────────────────────────────────────────────

app = FastAPI(title="Harbore Report Engine", version="1.0.0")

@app.get("/health")
def health():
    return {"status": "ok", "service": "harbore-report"}

@app.post("/report/generate")
async def generate_report(req: ReportRequest):
    if req.format == "pdf":
        buf = generate_pdf(req)
        media_type = "application/pdf"
        filename = f"harbore-{req.scan_id[:8]}.pdf"
    else:
        buf = generate_docx(req)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"harbore-{req.scan_id[:8]}.docx"

    return StreamingResponse(
        io.BytesIO(buf),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# ─── Word (.docx) generator ───────────────────────────────────────────────────

def generate_docx(req: ReportRequest) -> bytes:
    doc = Document()
    _set_docx_margins(doc)
    _setup_docx_styles(doc)

    # ── Cover page ────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Harbore")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)

    sub = doc.add_paragraph()
    sub.add_run("API Security Assessment Report").font.size = Pt(18)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.add_run(f"Scan: {req.scan_name}\n")
    meta.add_run(f"Date: {datetime.now().strftime('%B %d, %Y')}\n")
    meta.add_run(f"Targets assessed: {req.target_count}\n")
    meta.add_run(f"Scan ID: {req.scan_id}\n")
    if req.duration_mins:
        meta.add_run(f"Duration: {req.duration_mins:.0f} minutes\n")

    doc.add_page_break()

    # ── Executive Summary ─────────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    if req.executive_summary:
        doc.add_paragraph(req.executive_summary)
    else:
        stats_str = " | ".join(f"{k.title()}: {v}" for k, v in req.stats.items() if v > 0)
        doc.add_paragraph(
            f"A security assessment was conducted against {req.target_count} API targets. "
            f"Finding summary: {stats_str or 'No findings'}."
        )

    # Findings summary table
    doc.add_heading("Finding Summary", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Count"
    for sev in SEV_ORDER:
        count = req.stats.get(sev, 0)
        if count > 0:
            row = table.add_row().cells
            row[0].text = sev.capitalize()
            row[1].text = str(count)
            _color_cell(row[0], SEV_COLORS_RGB[sev])

    doc.add_page_break()

    # ── Findings by severity ───────────────────────────────────────────────────
    doc.add_heading("Findings", level=1)
    if req.critical_section:
        doc.add_paragraph(req.critical_section)

    for sev in SEV_ORDER:
        sev_findings = [f for f in req.findings if f.severity == sev]
        if not sev_findings:
            continue

        h = doc.add_heading(f"{sev.upper()} Severity ({len(sev_findings)})", level=2)
        h.runs[0].font.color.rgb = SEV_COLORS_RGB[sev]

        for i, f in enumerate(sev_findings, 1):
            doc.add_heading(f"{i}. {f.title}", level=3)

            meta_table = doc.add_table(rows=0, cols=2)
            meta_table.style = "Table Grid"
            _add_meta_row(meta_table, "Severity",   f.severity.capitalize())
            _add_meta_row(meta_table, "CVSS Score",  f"{f.cvss_score:.1f}" if f.cvss_score else "N/A")
            _add_meta_row(meta_table, "Module",      f.module)
            if f.endpoint:
                _add_meta_row(meta_table, "Endpoint", f"{f.method or 'GET'} {f.endpoint}")
            if f.owasp_ref:
                _add_meta_row(meta_table, "OWASP", f.owasp_ref)
            if f.pci_requirement:
                _add_meta_row(meta_table, "PCI DSS", f.pci_requirement)
            if f.cwe_id:
                _add_meta_row(meta_table, "CWE", f.cwe_id)

            doc.add_paragraph()

            doc.add_heading("Description", level=4)
            doc.add_paragraph(f.ai_summary or f.description)

            if f.ai_remediation:
                doc.add_heading("Remediation", level=4)
                doc.add_paragraph(f.ai_remediation)

            if f.request:
                doc.add_heading("Evidence — Request", level=4)
                p = doc.add_paragraph(f.request[:2000])
                p.style = doc.styles["No Spacing"]
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(8)

            if f.response:
                doc.add_heading("Evidence — Response", level=4)
                p = doc.add_paragraph(f.response[:2000])
                p.style = doc.styles["No Spacing"]
                p.runs[0].font.name = "Courier New"
                p.runs[0].font.size = Pt(8)

            doc.add_paragraph()

    doc.add_page_break()

    # ── Remediation roadmap ───────────────────────────────────────────────────
    doc.add_heading("Remediation Roadmap", level=1)
    if req.remediation_priorities:
        doc.add_paragraph(req.remediation_priorities)

    # ── PCI DSS section ───────────────────────────────────────────────────────
    pci_findings = [f for f in req.findings if f.pci_requirement]
    if pci_findings:
        doc.add_page_break()
        doc.add_heading("PCI DSS Compliance", level=1)
        if req.pci_narrative:
            doc.add_paragraph(req.pci_narrative)
        doc.add_heading(f"PCI DSS Findings ({len(pci_findings)})", level=2)
        for f in pci_findings:
            doc.add_heading(f.title, level=3)
            p = doc.add_paragraph()
            p.add_run(f"Requirement: {f.pci_requirement}\n").bold = True
            p.add_run(f.description)

    # ── Failure log ───────────────────────────────────────────────────────────
    if req.failures:
        doc.add_page_break()
        doc.add_heading("Scan Failures", level=1)
        doc.add_paragraph(
            f"{len(req.failures)} scan job(s) failed after maximum retry attempts. "
            "These targets may require manual testing."
        )
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Target", "Module", "Attempts", "Error"]):
            hdr[i].text = h
        for fail in req.failures:
            row = table.add_row().cells
            row[0].text = fail.target[:60]
            row[1].text = fail.module
            row[2].text = str(fail.attempts)
            row[3].text = fail.final_error[:100]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# ─── PDF generator ────────────────────────────────────────────────────────────

def generate_pdf(req: ReportRequest) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=20, spaceAfter=12, textColor=colors.HexColor("#3B82F6"))
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=14, spaceAfter=8, textColor=colors.HexColor("#E2E2E5"))
    h3 = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, spaceAfter=6, textColor=colors.HexColor("#9CA3AF"))
    body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, spaceAfter=6, leading=16)
    mono = ParagraphStyle("Mono", parent=styles["Code"], fontSize=8, spaceAfter=4, leading=12,
                          fontName="Courier", backColor=colors.HexColor("#1A1A1E"),
                          textColor=colors.HexColor("#9CA3AF"))

    # Cover
    story.append(Spacer(1, inch))
    story.append(Paragraph("Harbore", ParagraphStyle("Cover", fontSize=36, textColor=colors.HexColor("#3B82F6"), spaceAfter=12)))
    story.append(Paragraph("API Security Assessment Report", ParagraphStyle("CoverSub", fontSize=18, textColor=colors.HexColor("#E2E2E5"), spaceAfter=24)))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#374151")))
    story.append(Spacer(1, 12))
    story.append(Paragraph(f"<b>Scan:</b> {req.scan_name}", body))
    story.append(Paragraph(f"<b>Date:</b> {datetime.now().strftime('%B %d, %Y')}", body))
    story.append(Paragraph(f"<b>Targets:</b> {req.target_count}", body))
    story.append(Paragraph(f"<b>Scan ID:</b> {req.scan_id}", body))
    story.append(PageBreak())

    # Executive summary
    story.append(Paragraph("Executive Summary", h1))
    exec_text = req.executive_summary or f"Assessment of {req.target_count} API targets completed."
    story.append(Paragraph(exec_text, body))
    story.append(Spacer(1, 12))

    # Summary table
    table_data = [["Severity", "Count"]]
    for sev in SEV_ORDER:
        count = req.stats.get(sev, 0)
        if count > 0:
            table_data.append([sev.capitalize(), str(count)])

    if len(table_data) > 1:
        t = Table(table_data, colWidths=[2*inch, 1*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTSIZE",   (0,0), (-1,-1), 10),
            ("GRID",       (0,0), (-1,-1), 0.5, colors.HexColor("#374151")),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#111827"), colors.HexColor("#1F2937")]),
            ("TEXTCOLOR",  (0,1), (-1,-1), colors.HexColor("#E2E2E5")),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    story.append(PageBreak())

    # Findings
    story.append(Paragraph("Findings", h1))
    if req.critical_section:
        story.append(Paragraph(req.critical_section, body))

    for sev in SEV_ORDER:
        sev_findings = [f for f in req.findings if f.severity == sev]
        if not sev_findings:
            continue

        story.append(Paragraph(f"{sev.upper()} Severity ({len(sev_findings)})", h2))
        for i, f in enumerate(sev_findings, 1):
            story.append(Paragraph(f"{i}. {f.title}", h3))
            meta_rows = [
                ["Severity", f.severity.capitalize()],
                ["CVSS Score", f"{f.cvss_score:.1f}" if f.cvss_score else "N/A"],
                ["Module", f.module],
            ]
            if f.endpoint:  meta_rows.append(["Endpoint", f"{f.method or 'GET'} {f.endpoint[:80]}"])
            if f.owasp_ref: meta_rows.append(["OWASP", f.owasp_ref])
            if f.pci_requirement: meta_rows.append(["PCI DSS", f.pci_requirement])
            if f.cwe_id:    meta_rows.append(["CWE", f.cwe_id])

            mt = Table(meta_rows, colWidths=[1.5*inch, 4.5*inch])
            mt.setStyle(TableStyle([
                ("FONTSIZE", (0,0), (-1,-1), 9),
                ("TEXTCOLOR", (0,0), (0,-1), colors.HexColor("#9CA3AF")),
                ("TEXTCOLOR", (1,0), (1,-1), colors.HexColor("#E2E2E5")),
                ("GRID", (0,0), (-1,-1), 0.3, colors.HexColor("#374151")),
                ("BACKGROUND", (0,0), (-1,-1), colors.HexColor("#111827")),
                ("LEFTPADDING", (0,0), (-1,-1), 6),
            ]))
            story.append(mt)
            story.append(Spacer(1, 6))
            story.append(Paragraph(f.ai_summary or f.description, body))
            if f.ai_remediation:
                story.append(Paragraph("<b>Remediation:</b> " + f.ai_remediation, body))
            story.append(Spacer(1, 8))

    # Failure log
    if req.failures:
        story.append(PageBreak())
        story.append(Paragraph("Scan Failures", h1))
        fail_rows = [["Target", "Module", "Attempts", "Error"]]
        for fail in req.failures:
            fail_rows.append([fail.target[:40], fail.module, str(fail.attempts), fail.final_error[:50]])
        ft = Table(fail_rows, colWidths=[2*inch, 1*inch, 0.8*inch, 2.2*inch])
        ft.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#7F1D1D")),
            ("TEXTCOLOR",  (0,0), (-1,0), colors.white),
            ("FONTSIZE",   (0,0), (-1,-1), 8),
            ("GRID",       (0,0), (-1,-1), 0.3, colors.HexColor("#374151")),
            ("TEXTCOLOR",  (0,1), (-1,-1), colors.HexColor("#E2E2E5")),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.HexColor("#111827"), colors.HexColor("#1F2937")]),
        ]))
        story.append(ft)

    doc.build(story)
    return buf.getvalue()

# ─── Helpers ──────────────────────────────────────────────────────────────────

def _set_docx_margins(doc):
    from docx.oxml.ns import qn
    from docx.shared import Inches
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

def _setup_docx_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0xE2, 0xE2, 0xE5)

def _color_cell(cell, color: RGBColor):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def _add_meta_row(table, label: str, value: str):
    row = table.add_row().cells
    row[0].text = label
    row[1].text = value

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8091"))
    uvicorn.run("main:app", host="0.0.0.0", port=port, reload=False)


# ─── ISMS governance report (CB-Advisory template) ────────────────────────────
from isms_report import build_isms_docx


class Signatory(BaseModel):
    name: str = ""
    email: str = ""
    title: str = ""
    date: str = ""


class GovReportRequest(BaseModel):
    scan_id: str = ""
    document_title: str = "Security Assessment Report"
    subtitle: str = ""
    description: str = ""
    version: str = "0.1"
    effective_date: str = ""
    status: str = "Draft"
    owner: str = ""
    prepared: Signatory = Signatory()
    reviewed: Signatory = Signatory()
    approved: Signatory = Signatory()
    amendment_context: str = "Draft"
    amendment_revision: str = "First Draft"
    amendment_date: str = ""
    amendment_by: str = ""
    findings: list[FindingModel] = []


def _dump(m):
    return m.model_dump() if hasattr(m, "model_dump") else m.dict()


@app.post("/report/governance")
async def generate_governance(req: GovReportRequest):
    data = _dump(req)
    findings = data.pop("findings", [])
    buf = build_isms_docx(data, findings)
    filename = f"harbore-isms-{(req.scan_id or 'report')[:8]}.docx"
    return StreamingResponse(
        io.BytesIO(buf),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
