"""
Fills the CB-Advisory ISMS governance template (templates/isms_template.docx)
with user-supplied document metadata + signatories, and appends the scan
findings where the template says "START HERE…".

Kept dependency-light (python-docx only) so it can be unit-tested without the
rest of the report service.
"""
import io
import os
from collections import Counter

from docx import Document

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "templates", "isms_template.docx")
SEV_ORDER = ["critical", "high", "medium", "low", "info"]


def _set_para(p, text):
    """Replace a paragraph's text, keeping the first run's formatting."""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.text = text


def _set_cell(cell, text):
    _set_para(cell.paragraphs[0], text or "")


def build_isms_docx(data: dict, findings: list) -> bytes:
    def g(k, d=""):
        v = data.get(k)
        return v if v not in (None, "") else d

    doc = Document(TEMPLATE_PATH)

    # ── Header / metadata paragraphs (match by label) ──
    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "Document Title":
            _set_para(p, g("document_title", "Security Assessment Report"))
        elif t == "SubTitle":
            _set_para(p, g("subtitle", ""))
        elif t.startswith("Documentation Information:"):
            _set_para(p, "Documentation Information: " + g("description", ""))
        elif t.startswith("Version:"):
            _set_para(p, "Version: " + g("version", "0.1"))
        elif t.startswith("Effective Date:"):
            _set_para(p, "Effective Date: " + g("effective_date", "TBD"))
        elif t.startswith("Document Status:"):
            _set_para(p, "Document Status: " + g("status", "Draft"))
        elif t.startswith("Document Owner:"):
            _set_para(p, "Document Owner: " + g("owner", ""))
        elif t.startswith("START HERE"):
            _set_para(p, "Security Assessment Findings")

    # ── Signature table (table 0): rows 1/2/3 = Prepared/Reviewed/Approved ──
    if doc.tables:
        sig = doc.tables[0]
        for ri, key in {1: "prepared", 2: "reviewed", 3: "approved"}.items():
            if ri < len(sig.rows):
                s = data.get(key) or {}
                c = sig.rows[ri].cells  # cols: 0 label, 1 name, 2 email, 3 title, 4 date
                _set_cell(c[1], s.get("name", ""))
                _set_cell(c[2], s.get("email", ""))
                _set_cell(c[3], s.get("title", ""))
                _set_cell(c[4], s.get("date", ""))

    # ── Amendment record (table 1): first data row ──
    if len(doc.tables) > 1:
        amd = doc.tables[1]
        if len(amd.rows) > 1:
            c = amd.rows[1].cells
            _set_cell(c[0], g("version", "0.1"))
            _set_cell(c[1], g("amendment_context", "Draft"))
            _set_cell(c[2], g("amendment_revision", "First Draft"))
            _set_cell(c[3], g("amendment_date", g("effective_date", "")))
            _set_cell(c[4], g("amendment_by", (data.get("prepared") or {}).get("name", "")))

    # ── Findings (appended after "START HERE…", which is the final template line) ──
    counts = Counter((f.get("severity") or "info").lower() for f in findings)
    summary = ", ".join(f"{counts.get(s, 0)} {s}" for s in SEV_ORDER)
    doc.add_paragraph(f"{len(findings)} findings — {summary}.")

    if findings:
        table = doc.add_table(rows=1, cols=5)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for i, h in enumerate(["Severity", "Finding", "Endpoint", "CWE", "OWASP"]):
            _set_cell(table.rows[0].cells[i], h)
        rank = {s: i for i, s in enumerate(SEV_ORDER)}
        for f in sorted(findings, key=lambda x: rank.get((x.get("severity") or "info").lower(), 99))[:200]:
            c = table.add_row().cells
            _set_cell(c[0], (f.get("severity") or "").upper())
            _set_cell(c[1], f.get("title") or "")
            _set_cell(c[2], f.get("endpoint") or "")
            _set_cell(c[3], f.get("cwe_id") or "")
            _set_cell(c[4], f.get("owasp_ref") or "")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
